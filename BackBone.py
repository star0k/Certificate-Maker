import docx
import os
import pandas as pd
from docx2pdf import convert
from PyPDF2 import PdfFileMerger, PdfFileReader

StudentsData = []
files = []


def merge(files, out):
    mergedObject = PdfFileMerger()

    for file in files:
        mergedObject.append(PdfFileReader(file, 'rb'))
    mergedObject.write(f"{out}/All.pdf")

def getDataFromExcel (filepath) :
    print("getting data from excel...")
    xls = pd.ExcelFile(filepath)
    students = xls.sheet_names
    StudentsData = []

    for student in students :
            exl = pd.read_excel(xls,sheet_name=student)
            SName = student
            placer = 20
            sub = ''
            subjects = {}
            while not sub == 'TOTAL' :
                grades = {
                    'Passing': exl.iat[placer, 2],
                    'HomeWork' : exl.iat[placer,3],
                    'ClassWork' : exl.iat[placer,4],
                    'Exam' : exl.iat[placer,5],
                    'Total': exl.iat[placer, 6],
                    'Percentage': exl.iat[placer, 7],
                    'level': exl.iat[placer, 8],
                  }
                sub = exl.iat[placer,1]
                subjects[sub] = grades
                placer += 1
            Student = {}
            Student['Name'] = SName
            Student['marks'] = subjects
            per = "{:.2f}".format(exl.iat[33,2])
            Student['percentage'] = per
            Student['letter'] = exl.iat[33,3]
            Student['gpa'] = exl.iat[33,6]
            StudentsData.append(Student)
    return StudentsData
def rankStudents (StudentsData) :
    print("filling data..")
    forrank = {}
    for student in StudentsData :
        forrank[student['Name']] = student['percentage']
    forrank = {k: v for k, v in sorted(forrank.items(), reverse=True , key=lambda item: item[1])}
    frank = []
    for key in forrank.keys() :
        frank.append(key)
    rank = 1
    for student in StudentsData :
        student['Rank'] = frank.index(student['Name']) + 1
    return StudentsData


def fillWord(tamplate, StudentsData, Year, Term, Sclass ):
    # ROWS HEADERS
    headplace = {
        'HOMEWORK': 3,
        'CLASSWORK': 4,
        'EXAM': 5,
        'TOTAL': 6,
        'PERC': 2,
        'GRADE': 7,
        'letter': 8,
    }

    # COL SUBJECTS
    subcols = {
        'Math': 3,
        'English': 4,
        'Arabic': 5,
        'Science': 6,
        'Physics': 7,
        'ICT + Robot': 8,
        'İslamic': 9,
        'Physical Education': 10,
        'Turkish': 11,
        'Art': 12,
        'Attitude & Behaviour': 13,
        'TOTAL': 14,
        'percentage': 15,
        'gpa': 15,
        'Rank': 16,
        'letter': 15,
    }

    document = docx.Document(tamplate)
    NameTable = document.tables[1]
    GradesTable = document.tables[2]
    files = []
    for student in StudentsData:
        # Set Name
        NameTable.rows[0].cells[0].text = student['Name']
        # Set Class
        NameTable.rows[1].cells[0].text = Sclass
        # Set Term
        GradesTable.rows[0].cells[4].text = GradesTable.rows[0].cells[4].text.replace('{{termS}}', Term)
        # Set Year
        GradesTable.rows[0].cells[8].text = GradesTable.rows[0].cells[8].text.replace('{{year}}', Year)
        # Set Rank
        GradesTable.rows[16].cells[6].text = str(student['Rank'])
        # GPA
        GradesTable.rows[15].cells[6].text = str(student['gpa'])
        # set percentage
        GradesTable.rows[15].cells[headplace['PERC']].text = str(student['percentage'])
        # set letter
        GradesTable.rows[15].cells[3].text = str(student['letter'])

        # grades rows
        for subject in student['marks']:
            GradesTable.rows[subcols[subject]].cells[headplace['HOMEWORK']].text = str(
                student['marks'][subject]['HomeWork'])
            GradesTable.rows[subcols[subject]].cells[headplace['CLASSWORK']].text = str(
                student['marks'][subject]['ClassWork'])
            GradesTable.rows[subcols[subject]].cells[headplace['EXAM']].text = str(student['marks'][subject]['Exam'])
            GradesTable.rows[subcols[subject]].cells[headplace['TOTAL']].text = str(student['marks'][subject]['Total'])
            GradesTable.rows[subcols[subject]].cells[headplace['GRADE']].text = str(
                student['marks'][subject]['Percentage'])
            GradesTable.rows[subcols[subject]].cells[headplace['letter']].text = str(student['marks'][subject]['level'])

        try:
            os.mkdir(os.path.abspath(f"Grade {Sclass}"))
            os.mkdir(os.path.abspath(f"Grade {Sclass}/files"))
            os.mkdir(os.path.abspath(f"Grade {Sclass}/files/word"))
            os.mkdir(os.path.abspath(f"Grade {Sclass}/files/pdfs"))
        except:
            pass
        filename = f"Grade {Sclass}/files/word/{student['Name']}.docx"
        document.save(f"{filename}")
        print(f"{student['Name']} is done , next....")
        files.append(filename)
    print('word files done.')
    return files
def convert2PDF (Sclass) :
    print("converting files..")
    try :
        convert(f"Grade {Sclass}/files/word",f"Grade {Sclass}/files/pdfs/")
    except :
        print ("ok")
        pass
def mergeFiles (Sclass,files) :
    print("merging files..")

    PDFList = [file.replace('docx','pdf') for file in files ]
    PDFList = [file.replace('word','pdfs') for file in PDFList]
    merge(PDFList , f"Grade {Sclass}")
    return os.path.abspath(f"Grade {Sclass}/All.pdf")
